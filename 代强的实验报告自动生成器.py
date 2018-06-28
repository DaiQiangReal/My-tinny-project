import docx
import os
def sor(name):
	if name[1]=='.':
		return name[0]
	else:
		return name[0:2]

def picname(name):
	if name[3]=='.':
		return name[0:3]
	else:
		return name[0:4]

def main():
	print("			代强的实验报告自动生成器\n代码格式为1.1.c 截图格式为1.1.PNG\n")
	os.system("title 代强的实验报告自动生成器")
	os.system("color 2e")
	belong='1'
	pages=docx.Document()
	path=input("输入你的代码的路径\n")
	picpath=input("输入你截图的路径\n")
	anspath=input("输入需要存放实验报告的地方\n")
	change={'1':'实验一','2':'实验二','3':'实验三','4':'实验四','5':'实验五','6':'实验六',
	'7':'实验七','8':'实验八','9':'实验九','10':'实验十','11':'实验十一','12':'实验十二',
	'13':'实验十三','14':'实验十四','15':'实验十五','16':'实验十六'}
	for filepath, dirs, files in os.walk(path):
		for some in files:
			f=open(path+"\\"+str(some),"r",encoding="UTF-8")
			text=f.read()
			pages.add_paragraph(str(some))
			pages.add_paragraph(text)
			pages.add_picture(picpath+'\\'+picname(some)+'.PNG')
			pages.add_paragraph("******************************************************************\n")
			if files.index(some)+1==len(files):
				pages.save(anspath+"\\"+change[sor(files[files.index(some)])]+".docx")
				break
			if sor(files[files.index(some)+1])!=belong:
				pages.save(anspath+"\\"+change[sor(files[files.index(some)])]+".docx")
				pages=docx.Document()
				belong=sor(files[files.index(some)+1])

if __name__ == '__main__':
    main()
